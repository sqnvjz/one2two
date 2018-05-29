import os
import pandas as pd

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from utils import convert

root_path = os.path.dirname(os.path.abspath(__file__))


def read_xls(file_path):
    df = pd.read_excel(file_path)
    return df


def read_template_nl(word_path):
    dc = Document(word_path)
    parag_1 = dc.paragraphs[1]
    parag_2 = dc.paragraphs[2]
    parag_3 = dc.paragraphs[3]
    return {
        'parag_1': parag_1,
        'parag_2': parag_2,
        'parag_3': parag_3
    }


def write_word_nl(template, dst_path, data_dict):
    dc = Document()
    parag_1 = template['parag_1']
    text_1 = parag_1.text
    text_1_template = text_1.replace('01', '{}').replace('卓振兴', '{}').replace('15860747547', '{}')
    for i, d in enumerate(data_dict):
        tel = d['用户名']
        name = d['真实姓名']
        no = str(i + 1) if (i + 1) > 99 else str(i + 1).zfill(2)
        text_1 = text_1_template.format(no, name, tel)
        dc.add_paragraph(text_1, parag_1.style)
        parag_2 = template['parag_2']
        dc.add_paragraph(parag_2.text, parag_2.style)
        parag_3 = template['parag_3']
        dc.add_paragraph(parag_3.text, parag_3.style)
        dc.add_paragraph(parag_3.text, parag_3.style)
    dc.save(dst_path)


def read_template_info(word_path):
    dc = Document(word_path)
    header = dc.paragraphs[0]
    number = dc.paragraphs[1]
    blank = dc.paragraphs[2]
    table_1 = dc.tables[0]
    return {
        'table_1': table_1,
        'header': header,
        'number': number,
        'blank': blank
    }


def write_word_info(template, dst_path, data_dict, prefix='A3', start=1):
    dc = Document()
    table_1 = template['table_1']
    header = template['header']
    number = template['number']
    blank = template['blank']

    rows = len(table_1.rows)
    cols = len(table_1.columns)
    style = table_1.style

    for i, d in enumerate(data_dict):
        no = prefix + str(start + i).zfill(4)
        name = d['真实姓名']
        tel = d['用户名']
        money = d['商户预定PV消费量'] * 7
        money_ch = convert(str(money))

        p_header = dc.add_paragraph('信息表', header.style)
        p_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_header.runs[0].font.size = Pt(26)
        p_header.runs[0].font.name = 'Times'
        p_header.runs[0].font.bold = True


        p_number = dc.add_paragraph('编号：' + no, number.style)
        p_number.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_number.runs[0].font.size = Pt(16)
        p_number.runs[0].font.name = 'Times(宋体)'
        p_number.runs[0].font.bold = True

        tb = dc.add_table(rows=rows, cols=cols, style=style)
        tb.style.font.size = Pt(14)
        tb.cell(0, 0).text = table_1.cell(0, 0).text
        tb.cell(0, 1).text = name
        tb.cell(0, 2).text = table_1.cell(0, 2).text
        tb.cell(0, 3).text = table_1.cell(0, 3).text
        tb.cell(1, 0).text = table_1.cell(1, 0).text
        tb.cell(1, 1).text = str(tel)
        tb.cell(1, 1).merge(tb.cell(1, 2))
        tb.cell(1, 1).merge(tb.cell(1, 3))
        tb.cell(2, 0).text = table_1.cell(2, 0).text
        tb.cell(2, 1).text = '大写 ' + money_ch + '整'
        tb.cell(2, 1).merge(tb.cell(2, 2))
        tb.cell(2, 3).text = '小写 ' + str(money)
        tb.cell(3, 0).text = table_1.cell(3, 0).text
        tb.cell(3, 1).text = table_1.cell(3, 1).text
        tb.cell(3, 2).text = table_1.cell(3, 2).text
        tb.cell(3, 3).text = table_1.cell(3, 3).text
        for idx in range(4, 10):
            tb.cell(idx, 0).text = table_1.cell(idx, 0).text
            tb.cell(idx, 1).text = table_1.cell(idx, 1).text
            tb.cell(idx, 1).merge(tb.cell(idx, 2))
            tb.cell(idx, 1).merge(tb.cell(idx, 3))
        for _ in range(0, 14):
            dc.add_paragraph(blank.text, blank.style)

    dc.save(dst_path)
    pass


if __name__ == '__main__':
    # print('启动文件转换程序...')
    # prefix = input('请输入编号，例如 A3：')
    # start = input('请输入开始序号，例如 1：')
    #
    # print('回车后开始执行程序...')
    # start = int(start)
    prefix = 'A4'
    start = 1

    df = read_xls(os.path.join(root_path, 'files/source.xlsx'))
    data_dict = df.to_dict('records')

    tp_dict = read_template_nl(os.path.join(root_path, 'files/namelist.docx'))
    write_word_nl(tp_dict, os.path.join(root_path, 'dist/namelist_out.docx'), data_dict)

    tp_dict = read_template_info(os.path.join(root_path, 'files/info.docx'))
    write_word_info(tp_dict, os.path.join(root_path, 'dist/info_out.docx'), data_dict, prefix, start)

    print('end')
